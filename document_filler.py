import datetime
import os
import json
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.dropdown import DropDown
from kivy.uix.scrollview import ScrollView
from kivy.core.window import Window
from kivy.uix.popup import Popup
from kivy.uix.spinner import Spinner
from kivy.uix.widget import Widget
from kivy.graphics import Color, RoundedRectangle
from kivy.uix.filechooser import FileChooserIconView
from kivy.clock import Clock



from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import logging
from kivy.logger import Logger
import sys
import platform
import subprocess

# Configure logging
logging.basicConfig(level=logging.DEBUG)
Logger.info("Application started")

# Set window size for a more consistent UI look
Window.size = (600, 800)

# Function to handle file paths in both dev and packaged modes
def resource_path(relative_path):
    """ Get the absolute path to the resource, works for both dev and PyInstaller """
    try:
        base_path = sys._MEIPASS  # PyInstaller stores files in this temporary folder
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Function to handle writing to a user-writable directory (e.g., the user's home directory)
def writable_path(filename):
    """ Returns a writable path for the file in the user's home directory """
    return os.path.join(os.path.expanduser("~"), filename)

# File to store dropdown options (read from the app bundle, write to the home directory)
OPTIONS_FILE = writable_path('dropdown_options.json')

# Function to load dropdown options
def load_options():
    # If the file exists in the writable location, load it from there
    if os.path.exists(OPTIONS_FILE):
        with open(OPTIONS_FILE, 'r') as f:
            return json.load(f)
    else:
        # If the file doesn't exist in the writable location, copy it from the app resources
        default_options = {
            'NAME_OPTIONS': ["Fannan Mhamed", "El Harbouj Mohammed", "Hassan Laarbi"],
            'CONS_ID_OPTIONS': ["4966-7777-78", "4025-2450-53", "4028-5586-53"],
            'FLIGHT_OPTIONS': ["AT201", "AT203"],
            'TAG_OPTIONS': ["SPSM (Green label)", "Serveurs (Yellow label)", "1005 (white label)"],
            'DESC_OPTIONS': [f"Diplomatic box# {j}" for j in list(range(1, 11)) + list(range(101, 111))]
        }
        # Save default options to the writable location for future use
        save_options(default_options)
        return default_options

# Function to save dropdown options to the writable location
def save_options(options):
    with open(OPTIONS_FILE, 'w') as f:
        json.dump(options, f, indent=4)

# Load options at the start
OPTIONS = load_options()

# Function that fills the placeholders in the document
def fill_placeholders(doc_path, output_path, word_list, self):
    doc = Document(doc_path)
    placeholder_format = "{placeholder}"
    boxes_index = 3  # Now the number of boxes is in index 3

    # Try to extract the number of boxes
    try:
        num_boxes = int(word_list[boxes_index])  # Convert the number of boxes to an integer
    except ValueError:
        raise ValueError(f"Expected an integer for number of boxes, but got '{word_list[boxes_index]}'")

    # Iterate through paragraphs in the document and replace placeholders
    word_index = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder_format.format(placeholder=word_index + 1) in run.text and word_index < len(word_list):
                if word_index == boxes_index:
                    box_text = f"{num_boxes} diplomatic box" if num_boxes == 1 else f"{num_boxes} diplomatic boxes"
                    run.text = run.text.replace(placeholder_format.format(placeholder=word_index + 1), box_text)
                else:
                    run.text = run.text.replace(placeholder_format.format(placeholder=word_index + 1), word_list[word_index])
                run.font.color.rgb = RGBColor(0, 0, 0)
                word_index += 1

    if doc.tables == 0:
        self.result_label.text = "Your document doesn't contain any tables."
        return

    table = doc.tables[0]
    tag_start_index = 5 
    # Assuming the table is the first table in the document
    for i in range(num_boxes):
        row = table.add_row()


        # Set the row height to exactly 0.32 inches
        tr = row._tr  # Access the row's XML element
        trPr = tr.get_or_add_trPr()  # Get or create the row properties element
        trHeight = OxmlElement('w:trHeight')  # Create a row height element
        trHeight.set(qn('w:val'), str(int(Inches(0.000509090909091))))  # Convert 0.32 inches to twips (1 inch = 1440 twips)
        trHeight.set(qn('w:hRule'), 'exact')  # Set the height rule to 'exact'
        trPr.append(trHeight)  # Append the height element to the row properties

        # Add the tag to the first column (left cell)
        cell_1 = row.cells[1]
        tag_value = word_list[tag_start_index + 2 * i]  # Every 2nd item is a tag
        paragraph_1 = cell_1.paragraphs[0]
        run_1 = paragraph_1.add_run(f"{tag_value}")
        run_1.font.size = Pt(14)
        run_1.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
        paragraph_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add the description to the second column (right cell)
        cell_2 = row.cells[0]
        description_value = word_list[tag_start_index + 2 * i + 1]  # Every other item is a description
        paragraph_2 = cell_2.paragraphs[0]
        run_2 = paragraph_2.add_run(f"{description_value}")
        run_2.font.size = Pt(14)
        run_2.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
        paragraph_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save the document
    doc.save(output_path)


class ProfessionalApp(App):
    def build(self):
        self.input_doc = None
        scroll_view = ScrollView(size_hint=(1, None), size=(Window.width, Window.height))
        self.options = OPTIONS   # Store options in the a
        Window.bind(on_drop_file=self._on_file_drop)

        self.main_layout = BoxLayout(orientation='vertical', padding=30, spacing=20, size_hint_y=None)
        self.main_layout.bind(minimum_height=self.main_layout.setter('height'))

        # Title
        title_label = Label(text="ENVOI PREMIER RL", font_size=100, size_hint_y=None, height=70, bold=True)
        self.main_layout.add_widget(title_label)

        # Underline
        underline = Widget(size_hint_y=None, height=2)
        with underline.canvas:
            Color(0, 0, 0, 1)
            RoundedRectangle(pos=underline.pos, size=(Window.width - 60, 2))
        self.main_layout.add_widget(underline)

        # Name field
        self.name_input = self.create_centered_input("Enter or Select Name")
        self.name_input.bind(focus=self.show_name_dropdown)
        self.main_layout.add_widget(Label(text="Name:", size_hint_y=None, height=30))
        self.main_layout.add_widget(self.name_input)

        # Consular ID field
        self.cons_id_input = self.create_centered_input("Enter or Select Consular ID")
        self.cons_id_input.bind(focus=self.show_cons_id_dropdown)
        self.main_layout.add_widget(Label(text="Consular ID:", size_hint_y=None, height=30))
        self.main_layout.add_widget(self.cons_id_input)

        # Number of boxes
        self.box_count_input = self.create_centered_input("Select Number of Boxes", readonly=True)
        self.box_count_input.bind(focus=self.show_box_count_dropdown)
        self.main_layout.add_widget(Label(text="Number of Boxes:", size_hint_y=None, height=30))
        self.main_layout.add_widget(self.box_count_input)

        # Flight Number field (properly restored)
        self.flight_input = self.create_centered_input("Enter or Select Flight Number")
        self.flight_input.bind(focus=self.show_flight_dropdown)
        self.main_layout.add_widget(Label(text="Flight Number:", size_hint_y=None, height=30))
        self.main_layout.add_widget(self.flight_input)

        # Box Details Label (initially hidden until number of boxes is selected)
        self.box_details_label = Label(text="Box Details:", size_hint_y=None, height=30)
        self.box_details_label.opacity = 0  # Initially hidden
        self.main_layout.add_widget(self.box_details_label)

        # Dynamic inputs container for tag/description pairs
        self.dynamic_inputs_container = BoxLayout(orientation='vertical', spacing=10, size_hint_y=None)
        self.dynamic_inputs_container.bind(minimum_height=self.dynamic_inputs_container.setter('height'))
        self.main_layout.add_widget(self.dynamic_inputs_container)

        # Add a button to upload document with heartbeat effect
        self.upload_doc_button = Button(
            text="drag and drop document (.docx)",
            size_hint=(1, None), height=50,
            background_color=(1, 0, 0, 1)  # Initial color (red)
        )
        self.upload_doc_button.bind(on_release=self.show_file_chooser)  # Bind to show file chooser
        self.main_layout.add_widget(self.upload_doc_button)

        # Start the heartbeat effect
        self.document_selected = False
        self.heartbeat_event = Clock.schedule_interval(self.heartbeat_effect, 0.5) 

        # Submit button
        self.submit_button = Button(text="Submit", size_hint=(None, None), size=(200, 50), pos_hint={'center_x': 0.5}, background_normal='', background_color=(0.3, 0.5, 0.7, 1), color=(1, 1, 1, 1))
        self.submit_button.bind(on_press=self.process_document)
        self.main_layout.add_widget(Widget(size_hint_y=None, height=20))
        self.main_layout.add_widget(self.submit_button)

        self.result_label = Label(text="", size_hint_y=None, height=50)
        self.main_layout.add_widget(self.result_label)

        # Add Item Button (dimmed appearance)
        add_item_button = Button(
            text="Update dropdowns",
            size_hint=(None, None),
            size=(270, 50),
            pos_hint={'center_x': 0.5},
            background_normal='',  # Make the button dimmed
            background_color=(0.5, 0.5, 0.5, 1),  # Dimmed color (gray)
            color=(1, 1, 1, 1)  # Text color remains white
        )

        # Bind the button to show the password popup when clicked
        add_item_button.bind(on_press=self.show_password_popup)

        # Add the button to the layout
        self.main_layout.add_widget(add_item_button)

        scroll_view.add_widget(self.main_layout)
        return scroll_view
    
    def show_add_item_popup(self, instance):
        content = BoxLayout(orientation='vertical', padding=10, spacing=10)
        
        dropdown_selector = Spinner(
            text='Select Dropdown',
            values=('Name', 'Consular ID', 'Flight', 'Tag', 'Description'),
            size_hint_y=None,
            height=44
        )
        
        new_item_input = TextInput(
            hint_text='Enter new item',
            multiline=False,
            size_hint_y=None,
            height=44
        )
        
        add_button = Button(
            text='Add Item',
            size_hint_y=None,
            height=44
        )
        
        content.add_widget(dropdown_selector)
        content.add_widget(new_item_input)
        content.add_widget(add_button)
        
        popup = Popup(
            title='Add New Item',
            content=content,
            size_hint=(None, None),
            size=(400, 300)
        )
        
        def check_dropdown_selection(*args):
            if dropdown_selector.text == 'Select Dropdown':  # Nothing selected
                self.start_heartbeat_effect(dropdown_selector)  # Start heartbeat if no selection
            else:
                self.add_new_item(dropdown_selector.text, new_item_input.text, popup)
        
        add_button.bind(on_press=check_dropdown_selection)
        
        popup.open()

    def start_heartbeat_effect(self, widget):
        # Define a counter to keep track of the number of heartbeats
        self.heartbeat_count = 0
        
        def beat(dt):
            if self.heartbeat_count < 8:  # 6 toggles (3 red beats)
                current_color = widget.background_color
                if current_color == [1, 0.2, 0.2, 1]:  # If red
                    widget.background_color = (1, 1, 1, 1)  # Toggle to white
                else:
                    widget.background_color = (1, 0.2, 0.2, 1)  # Toggle to red
                self.heartbeat_count += 1
            else:
                widget.background_color = (1, 1, 1, 1)  # Reset to white after heartbeat
                Clock.unschedule(heartbeat_event)  # Stop the heartbeat effect
        
        # Schedule the heartbeat effect to run every 0.2 seconds
        heartbeat_event = Clock.schedule_interval(beat, 0.1)

    def add_new_item(self, dropdown_name, new_item, popup):
        if not new_item.strip():
            return  # Don't add empty items

        option_key = f'{dropdown_name.upper().replace(" ", "_")}_OPTIONS'
        if option_key in self.options:
            if new_item not in self.options[option_key]:
                self.options[option_key].append(new_item)
                save_options(self.options)  # Save updated options to file
                self.update_dropdowns()  # Ensure dropdowns are refreshed in the UI

        popup.dismiss()

    def update_dropdowns(self):
        # Update name dropdown
        self.name_input.bind(focus=self.show_name_dropdown)

        # Update consular ID dropdown
        self.cons_id_input.bind(focus=self.show_cons_id_dropdown)

        # Update flight dropdown
        self.flight_input.bind(focus=self.show_flight_dropdown)

        # Update tag and description dropdowns in dynamic inputs
        for box in self.dynamic_inputs_container.children:
            tag_input, desc_input = box.children
            tag_input.bind(focus=self.create_dropdown_function(tag_input, self.options['TAG_OPTIONS']))
            desc_input.bind(focus=self.create_dropdown_function(desc_input, self.options['DESC_OPTIONS']))

    def create_dropdown_function(self, text_input, options):
        def show_dropdown(instance, value):
            if value:
                dropdown = DropDown()
                for option in options:
                    btn = Button(text=str(option), size_hint_y=None, height=44)
                    btn.bind(on_release=lambda btn: self.select_from_dropdown(btn, text_input, dropdown))
                    dropdown.add_widget(btn)
                dropdown.open(text_input)
        return show_dropdown

    def create_centered_input(self, hint_text, readonly=False):
        return TextInput(hint_text=hint_text, size_hint=(None, None), size=(540, 50), multiline=False, readonly=readonly, pos_hint={'center_x': 0.5})

    def show_box_count_dropdown(self, instance, value):
        if value:
            dropdown = DropDown()
            for i in range(1, 7):
                btn = Button(text=str(i), size_hint_y=None, height=44)
                btn.bind(on_release=lambda btn: self.select_box_count(btn.text, dropdown))
                dropdown.add_widget(btn)
            dropdown.open(instance)

    def select_box_count(self, text, dropdown):
        self.box_count_input.text = text
        dropdown.dismiss()
        self.create_dynamic_dropdowns(int(text))

    def create_dynamic_dropdowns(self, num_boxes):
        self.dynamic_inputs_container.clear_widgets()
        self.box_details_label.opacity = 1

        for i in range(num_boxes):
            box = BoxLayout(orientation='horizontal', spacing=10, size_hint_y=None, height=60)
            tag_input = self.create_centered_input(f"Select Tag {i+1}")
            tag_input.bind(focus=self.create_dropdown_function(tag_input, self.options['TAG_OPTIONS']))
            desc_input = self.create_centered_input(f"Select Box Description {i+1}")
            desc_input.bind(focus=self.create_dropdown_function(desc_input, self.options['DESC_OPTIONS']))
            box.add_widget(tag_input)
            box.add_widget(desc_input)
            self.dynamic_inputs_container.add_widget(box)

    def select_from_dropdown(self, button, text_input, dropdown):
        text_input.text = button.text
        dropdown.dismiss()

    def show_name_dropdown(self, instance, value):
        if value:
            self.show_generic_dropdown(instance, self.options['NAME_OPTIONS'], self.name_input)

    def show_cons_id_dropdown(self, instance, value):
        if value:
            self.show_generic_dropdown(instance, self.options['CONS_ID_OPTIONS'], self.cons_id_input)

    def show_flight_dropdown(self, instance, value):
        if value:
            self.show_generic_dropdown(instance, self.options['FLIGHT_OPTIONS'], self.flight_input)

    def show_generic_dropdown(self, instance, options, target_input):
        dropdown = DropDown()
        for option in options:
            btn = Button(text=option, size_hint_y=None, height=44)
            btn.bind(on_release=lambda btn: self.select_from_dropdown(btn, target_input, dropdown))
            dropdown.add_widget(btn)
        dropdown.open(instance)



    def show_file_chooser(self, instance):
        # Create a vertical layout for the popup content
        content = BoxLayout(orientation="vertical")
        
        # Create the FileChooserIconView to browse for .docx files
        filechooser = FileChooserIconView(
            filters=["*.docx"],  # Filter to show only .docx files
            path=os.path.expanduser("~")  # Start at the home directory
        )
        
        # Create 'Select' and 'Cancel' buttons
        button_layout = BoxLayout(orientation="horizontal", size_hint_y=None, height=50)
        select_button = Button(text="Select", size_hint=(1, None), height=40)
        cancel_button = Button(text="Cancel", size_hint=(1, None), height=40)
        
        # Add buttons to a horizontal layout for proper placement
        button_layout.add_widget(select_button)
        button_layout.add_widget(cancel_button)
        
        # Create the popup with the FileChooser and buttons
        popup = Popup(
            title="Select a DOCX file",
            content=content,
            size_hint=(0.9, 0.9)
        )
        
        # Bind the 'Select' button to handle file selection
        select_button.bind(on_release=lambda x: self.select_file(filechooser.selection, popup))
        cancel_button.bind(on_release=popup.dismiss)
        
        # Add file chooser and buttons to the content layout
        content.add_widget(filechooser)
        content.add_widget(button_layout)
        
        # Open the popup
        popup.open()

    def select_file(self, selection, popup):
        if selection:
            selected_file = selection[0]  # Get the first selected file
            if selected_file.lower().endswith('.docx'):
                # Update button text with the file name
                self.upload_doc_button.text = os.path.basename(selected_file)
                
                # Set the selected file as the input document for further processing
                self.input_doc = selected_file
                
                # Close the popup
                popup.dismiss()
                
                # Optionally log the selection
                Logger.info(f"Selected file: {selected_file}")
            else:
                Logger.warning("Invalid file type selected. Only .docx files are allowed.")
        else:
            Logger.warning("No file selected.")

        # Heartbeat effect to pulse the button's color
    def heartbeat_effect(self, dt):
        if not self.document_selected:
            current_color = self.upload_doc_button.background_color
            # Toggle between two shades of red for the heartbeat effect
            if current_color == [1, 0, 0, 1]:  # If red
                self.upload_doc_button.background_color = (1, 0.5, 0.5, 1)  # Light red
            else:
                self.upload_doc_button.background_color = (1, 0, 0, 1)
                



    def _on_file_drop(self, window, file_path, x, y):
        file_path_decoded = file_path.decode("utf-8")  # Decode the file path (it's a byte string)

        # Check if the file is a .docx
        if not file_path_decoded.lower().endswith('.docx'):
            # Notify user the document must be .docx
            self.result_label.text = "Invalid file type. Please select a .docx file."
            Logger.warning("Invalid file type dropped")
            
            # Reset the document state to allow for the next attempt
            self.document_selected = False
            self.input_doc = None  # Clear any previous input_doc
            return  # Discard the document and do nothing

        # Load the document
        try:
            doc = Document(file_path_decoded)
        except Exception as e:
            self.result_label.text = "Error reading the document."
            Logger.error(f"Error loading the .docx document: {e}")
            
            # Reset the document state to allow for the next attempt
            self.document_selected = False
            self.input_doc = None  # Clear any previous input_doc
            return

        # Iterate through paragraphs and remove trailing spaces and empty lines
        for paragraph in doc.paragraphs:
            # Remove trailing spaces in paragraph text
            paragraph.text = paragraph.text.rstrip()

        # Remove trailing empty paragraphs at the end
        # Ensure that the element is actually a child of the document element before removing it
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            try:
                # Check if the element is a valid child before attempting removal
                if doc.paragraphs[-1]._element in doc._element:
                    doc._element.remove(doc.paragraphs[-1]._element)
                else:
                    Logger.warning("Attempted to remove a paragraph element that is not a valid child.")
                    break  # If not a valid child, break the loop to prevent errors
            except Exception as e:
                Logger.error(f"Error removing paragraph element: {e}")
                break

        # If everything is valid, update global input_doc with resource_path
        self.input_doc = resource_path(file_path_decoded)

        # Extract and display only the file name, not the full path
        file_name = os.path.basename(file_path_decoded)
        self.upload_doc_button.text = file_name  # Update button text with file name

        self.document_selected = True  # Stop the heartbeat effect
        self.upload_doc_button.background_color = (0, 1, 0, 1)  # Change to green when valid file is dropped
        Clock.unschedule(self.heartbeat_event)  # Stop the pulsing effect


    def process_document(self, instance):
        Logger.info("Submit button pressed")
        
        # Check if document is selected
        if not self.input_doc:
            self.result_label.text = "Please select a document before submitting."
            Logger.warning("No document selected")
            print(self.input_doc)
            return

        # Check if all fields are filled
        if not self.all_fields_filled():
            self.result_label.text = "Please fill all fields before submitting."
            Logger.warning("Fields not filled")
            return

        # Get the input values
        name = self.name_input.text
        cons_id = self.cons_id_input.text
        box_count = self.box_count_input.text
        flight = self.flight_input.text

        # Prepare the word list for the document
        word_list = [name, cons_id, box_count, flight]

        # Collect tag and description dropdown values
        for box in self.dynamic_inputs_container.children:
            tag_input, desc_input = box.children
            word_list.append(tag_input.text)
            word_list.append(desc_input.text)

        # Set today's date as the first placeholder
        today_date = datetime.datetime.now().strftime("%B %d, %Y")
        word_list.insert(0, today_date)

        # Path to save the output document on the Desktop
        output_doc = os.path.join(os.path.expanduser("~"), "Desktop", "ENVOI PREMIER RL.docx")

        # Call the document-filling function with the selected document
        fill_placeholders(self.input_doc, output_doc, word_list, self)

        # Display confirmation
        self.result_label.text = "Document filled and saved to Desktop!"

        # Clear all input fields
        self.clear_all_fields()

        try:
            Logger.info(f"Document processing completed for: {name}")
            
            # Automatically open the document based on the OS
           # self.open_document(output_doc)
            
        except Exception as e:
            Logger.error(f"Error during document processing: {e}")
            raise

    # Function to open the document based on the OS
    def open_document(self, file_path):
        # Detect the platform and open the document accordingly
        try:
            if platform.system() == "Darwin":  # macOS
                subprocess.call(('open', file_path))
            elif platform.system() == "Windows":  # Windows
                os.startfile(file_path)
            elif platform.system() == "Linux":  # Linux
                subprocess.call(('xdg-open', file_path))
        except Exception as e:
            Logger.error(f"Failed to open document: {e}")
            self.result_label.text = f"Document saved to Desktop, but failed to open: {e}"


    def show_password_popup(self, instance):
            # Create a popup content layout
        content = BoxLayout(orientation='vertical', padding=10, spacing=10)

            # Password input field
        password_input = TextInput(
                hint_text='Enter password',
                multiline=False,
                password=True,  # Mask the input for password
                size_hint_y=None,
                height=44
        )

            # Submit button inside the popup
        submit_button = Button(
                text='Submit',
                size_hint_y=None,
                height=44
        )

            # Add the password input and submit button to the content layout
        content.add_widget(password_input)
        content.add_widget(submit_button)

            # Create the popup
        password_popup = Popup(
                title='Admin Access Required',
                content=content,
                size_hint=(None, None),
                size=(400, 200)
        )

        # Bind the submit button action
        submit_button.bind(on_press=lambda x: self.check_password(password_input.text, password_popup))

        # Open the popup
        password_popup.open()

    def check_password(self, entered_password, popup):
        # Check if the entered password is correct
        if entered_password == '1975':
            popup.dismiss()  # Close the popup
            self.show_add_item_popup(None)  # Proceed to the actual action (updating dropdowns)
        else:
            # Incorrect password - show an error
            error_popup = Popup(
                title='Error',
                content=Label(text='Incorrect password'),
                size_hint=(None, None),
                size=(300, 150)
            )
            error_popup.open()

    def all_fields_filled(self):
        if not all([self.name_input.text, self.cons_id_input.text, self.box_count_input.text, self.flight_input.text]):
            return False
        for box in self.dynamic_inputs_container.children:
            tag_input, desc_input = box.children
            if not tag_input.text or not desc_input.text:
                return False
        return True

    def clear_all_fields(self):
        self.name_input.text = ""
        self.cons_id_input.text = ""
        self.box_count_input.text = ""
        self.flight_input.text = ""
        self.dynamic_inputs_container.clear_widgets()


# Run the app
if __name__ == "__main__":
    ProfessionalApp().run()